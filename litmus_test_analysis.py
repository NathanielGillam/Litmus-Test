### Import supporting libraries
import psycopg2
import ssl
import getpass
import pytz
import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import docx
import io
import datetime
import scipy.ndimage as ndi
from scipy.interpolate import interp1d as intp
import skimage
from PIL import Image
from skimage.color import rgb2lab
from skimage.filters import threshold_otsu
from skimage.transform import hough_line, hough_line_peaks
from skimage.morphology import remove_small_holes, remove_small_objects
from skimage.measure import find_contours
from PIL import Image
from PIL.ExifTags import TAGS

### Configure Neon Connection String
NEON_CONN = "postgresql://neondb_owner:npg_RqKQsPe3U4HV@ep-billowing-waterfall-apt0oh58.c-7.us-east-1.aws.neon.tech/neondb?sslmode=require"

### User information
logged_in_user = getpass.getuser()
### Timestamp 
pst = pytz.timezone("America/Los_Angeles")

### Auto Connect to NEON
def get_connection():
    return psycopg2.connect(NEON_CONN, sslmode="require")

### DB initialization
def initialize_db():
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS litmus_results (
        id SERIAL PRIMARY KEY,
        timestamp TIMESTAMPTZ,
        spray_cell INTEGER,
        chemical_type TEXT,
        pass_fail TEXT,
        mean_width REAL,
        mean_deflection REAL,
        output_image BYTEA
        )
    """)
    conn.commit()
    cur.close()
    conn.close()

initialize_db()

### Configure streamlit webpage title, icon, and page width
st.set_option("client.toolbarMode", "minimal")
st.set_page_config(
    page_title = 'AO Litmus Test Analysis',
    page_icon = 'Active',
    layout = 'wide'
)

st.markdown(
    """
    <style>
    footer,
    .stDeployButton,
    .stToolbar {
        display: none;
    }
    </style>
    """,
    unsafe_allow_html=True
)


### Force the scrollbar to be a constant width. 
    # On some browsers, the scrollbar constantly switches back and forth between its thicker mouseover state
    # and its thinner non-mouseover state causing the page to resize cyclically.
st.markdown("""
                <html>
                    <head>
                    <style>
                        ::-webkit-scrollbar {
                            width: 5px;
                            }
                    </style>
                    </head>
                    <body>
                    </body>
                </html>
            """, unsafe_allow_html=True)

### Revise the size of the margin at the top of the page.
st.markdown("""
        <style>
               .block-container {
                    padding-top: 0.5rem;
        </style>
        """, unsafe_allow_html=True)
  
### Create the page sidebar with most of the user's controls
with st.sidebar:
### Some user inputs for DB reading
    st.subheader("Test Inputs")

    spray_cell = st.selectbox(
        "Spray Cell Number",
        options=["-- Select --", 1, 2, 3, 4],
        index=0
    )

    chemical_type = st.selectbox(
        "Chemical Type",
        options=["-- Select --", "M1b", "Glass", "MALP"],
        index=0
    )



    uploaded_image = st.file_uploader(label = '', type = ['jpg'], label_visibility = 'collapsed')
    st.divider()
    st.subheader('Image Settings')
    auto_crop = st.checkbox('Auto-crop (beta)', value = True)
    if not auto_crop:
        tc = st.number_input('Top edge crop', value = 25, step = 5)
        bc = st.number_input('Bottom edge crop', value = 1215, step = 5)
        lc = st.number_input('Left edge crop', value = 960, step = 5)
        rc = st.number_input('Right edge crop', value = 25, step = 5)
        crop_done = st.checkbox('Done Cropping', value = False)
    else:
        crop_done = False
    
    edit_center = st.checkbox('Adjust centerline identification', value = False)
    if edit_center:
        m_thresh = st.slider(label = 'Centerline threshold', min_value = 30, max_value = 60, value = 45)

    if auto_crop:
        st.markdown(body = "\n\n")
        st.markdown(body = """:grey[NOTE: The new auto-crop feature is still being tested and will likely not be completely reliable.
                    This feature is very sensitive to the condition of the edges of the litmus paper. Straight, clean cut edges will produce the best result.]""")

### Set the standard image resolution used to 200 dpi
image_dpi = 200

### Create a page header and seperate the page into columns for instructions, image results, and measurement results.      
st.header('Auto-Oxi Litmus Test Characterization')
inst_col, m1, crop_col = st.columns([0.28, 0.02, 0.7])
with crop_col:
    img_col, m2, val_col = st.columns([0.5, 0.01, 0.49])

tab1, tab2 = st.tabs(["Litmus Test", "Database Dashboard"])

with tab1:

    # Require user selections (but do NOT block the script)
    if spray_cell == "-- Select --" or chemical_type == "-- Select --":
        st.warning("Please select both Spray Cell Number and Chemical Type before uploading an image.")
        st.stop()   # stops only tab1, not tabs below

    # Instructions
    with inst_col:
        st.markdown(
            """Instructions:
            1. Scan litmus as a 200‑dpi JPG.
            2. Upload image in left sidebar.
            3. Crop or let Auto‑Crop run.
            4. Adjust centerline threshold if needed.
            5. Download the report to auto‑save to the database."""
        )
        st.divider()

    # UI placeholders
    with crop_col:
        crop_placeholder = st.empty()
    with img_col:
        image_placeholder = st.empty()
    with val_col:
        values_placeholder = st.empty()

    ############################################################
    # ------------------- IMAGE PROCESSING ---------------------
    ############################################################

    if uploaded_image is not None:

        # Load image
        image = skimage.io.imread(uploaded_image)
        meta_img = Image.open(uploaded_image)
        input_dpi = meta_img.info.get("dpi", (200,200))
        image_dpi = input_dpi[0]

        if image_dpi != 200:
            st.error(f"Input image resolution is {input_dpi}, expected 200 DPI.")

        # ------------------- AUTO CROP --------------------------
        if auto_crop:
            try:
                segments = skimage.segmentation.slic(image, n_segments=3, convert2lab=True)
                unborder = skimage.segmentation.clear_border(segments)
                mask = unborder > 0
                mask3 = np.stack([mask, mask, mask], axis=2)
                m_image = np.ma.masked_where(~mask3, image)

                tc = next(i for i in range(m_image.shape[0]) if not m_image.mask[i,:,:].all())
                bc_raw = next(i for i in range(m_image.shape[0]-1, -1, -1) if not m_image.mask[i,:,:].all())
                bc = m_image.shape[0] - bc_raw

                acceptance = 0.2
                for col in range(m_image.shape[1]):
                    if m_image.mask[tc:m_image.shape[0]-bc, col, :].mean() <= acceptance:
                        lc = col
                        break

                for col in range(m_image.shape[1]-1, -1, -1):
                    if m_image.mask[tc:m_image.shape[0]-bc, col, :].mean() <= acceptance:
                        rc = m_image.shape[1] - col
                        break

                crop_done = True

            except Exception:
                st.error("Auto‑crop failed — please crop manually.", icon="⚠️")
                crop_done = False

        # Manual cropping still supported
        cropped_image = skimage.util.crop(image, ((tc, bc), (lc, rc), (0,0)))

        # Crop preview
        crop_fig, ax = plt.subplots()
        if auto_crop:
            ax.imshow(skimage.segmentation.mark_boundaries(image, unborder))
        else:
            ax.imshow(image)
        ax.set_xticks([]); ax.set_yticks([])

        if not crop_done:
            crop_placeholder.write(crop_fig)

    ############################################################
    # ------------------- MEASUREMENT LOGIC --------------------
    ############################################################

    if uploaded_image is not None and crop_done:

        # LAB conversion + centerline threshold
        gray = rgb2lab(cropped_image)[:,:,0]
        g_thresh = m_thresh if edit_center else threshold_otsu(gray)

        mask = gray < g_thresh
        mask[:int(0.35*mask.shape[0])] = 0
        mask[int(0.65*mask.shape[0]):] = 0

        # Centerline via Hough transform
        h, theta, dist = hough_line(mask)
        _, best_angle, best_dist = hough_line_peaks(h, theta, dist, num_peaks=1)

        angle = best_angle[0]
        slope = np.tan(angle + np.pi/2)
        x0 = best_dist[0] * np.cos(angle)
        y0 = best_dist[0] * np.sin(angle)
        cntr = (x0, y0, slope)

        # Spray edge detection
        red_img = rgb2lab(cropped_image)[:,:,2]
        r_thresh = threshold_otsu(red_img)

        blr = ndi.uniform_filter(red_img, size=int(0.2*image_dpi))
        spray_mask = remove_small_holes(remove_small_objects(blr > r_thresh, 2000), 4000)

        contours = find_contours(spray_mask)
        characterize = len(contours) == 2

        ############################################################
        # ------------------ IF PATTERN DETECTED -------------------
        ############################################################

        if characterize:

            bottom = contours[1]
            top = contours[0]

            bot_x, bot_y = bottom[:,1], bottom[:,0]
            top_x, top_y = top[:,1], top[:,0]

            x = np.arange(cropped_image.shape[1])

            bot_interp = intp(bot_x, bot_y, fill_value="extrapolate")(x)
            top_interp = intp(top_x, top_y, fill_value="extrapolate")(x)
            centerline = cntr[2]*x + cntr[1]

            spray_width = abs(top_interp - bot_interp) / image_dpi
            top_height = (centerline - top_interp) / image_dpi
            bot_height = (bot_interp - centerline) / image_dpi
            apparent_center = (top_interp + bot_interp) / 2
            deflection = np.degrees(np.arctan((centerline - apparent_center) / image_dpi))

            # PASS/FAIL logic
            if np.isnan(spray_width.mean()):
                disp = "FAIL"
            elif spray_width.mean() < 1 or spray_width.mean() > 2:
                disp = "FAIL"
            elif abs(deflection.mean()) > 10:
                disp = "FAIL"
            elif (deflection.max() - deflection.min()) > 15:
                disp = "FAIL"
            elif abs(deflection).max() > 18:
                disp = "FAIL"
            else:
                disp = "PASS"

            # Show annotated image
            fig, ax = plt.subplots()
            ax.imshow(cropped_image)
            ax.plot(bot_x, bot_y, 'k')
            ax.plot(top_x, top_y, 'k')
            ax.axline((cntr[0], cntr[1]), slope=cntr[2], linestyle='--', color='white')
            ax.set_xticks([]); ax.set_yticks([])

            fig_mem = io.BytesIO()
            fig.savefig(fig_mem, bbox_inches='tight', pad_inches=0)

            # Results table
            results = pd.DataFrame({
                "mean":[spray_width.mean(), top_height.mean(), bot_height.mean(), deflection.mean()],
                "std":[spray_width.std(), top_height.std(), bot_height.std(), deflection.std()],
                "min":[spray_width.min(), top_height.min(), bot_height.min(), deflection.min()],
                "max":[spray_width.max(), top_height.max(), bot_height.max(), deflection.max()],
            }, index=[
                "Spray Width (in.)",
                "Top Width (in.)",
                "Bottom Width (in.)",
                "Deflection (deg.)"
            ])

            image_placeholder.write(fig)
            values_placeholder.dataframe(results)

        else:
            # → characterization failed
            disp = "FAIL"
            spray_width = None
            deflection = None

            fig_mem = io.BytesIO()
            fig, ax = plt.subplots()
            ax.imshow(cropped_image)
            ax.set_xticks([]); ax.set_yticks([])
            fig.savefig(fig_mem)

            st.error("Spray pattern edge detection failed.")

        ############################################################
        # ------------------ REPORT + SAVE TO DB ------------------
        ############################################################

        # Build DOCX report
        doc = docx.Document()
        now = datetime.datetime.now(pst)

        title = doc.add_paragraph()
        title_run = title.add_run("Litmus Test Characterization Report")
        title_run.bold = True
        title_run.font.size = docx.shared.Pt(18)

        doc.add_paragraph(now.strftime("%B %d, %Y - %I:%M %p"))
        pf = doc.add_paragraph().add_run(disp)
        pf.bold = True
        pf.font.size = docx.shared.Pt(14)

        # Add image
        fig_mem.seek(0)
        doc.add_picture(fig_mem, height=docx.shared.Inches(3))

        # Add table ONLY if characterize=True
        if characterize:
            table = doc.add_table(rows=results.shape[0]+1, cols=results.shape[1]+1)
            table.cell(0,0).text = ""

            for j,col in enumerate(results.columns):
                table.cell(0,j+1).text = col

            for i,row in enumerate(results.index):
                table.cell(i+1,0).text = row
                for j,col in enumerate(results.columns):
                    table.cell(i+1,j+1).text = f"{results.loc[row,col]:.3f}"

        # Final report bytes
        download_object = io.BytesIO()
        doc.save(download_object)

       # Prepare bytes for download
        download_object.seek(0)
        report_bytes = download_object.getvalue()
        
        if st.download_button(
            label="Download Report",
            data=report_bytes,
            file_name=f"litmus_{now.strftime('%m_%d_%y_%H_%M_%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            try:
                fig_mem.seek(0)
                image_bytes = fig_mem.getvalue()
        
                conn = get_connection()
                cur = conn.cursor()
        
                cur.execute("""
                    INSERT INTO litmus_results (
                        timestamp,
                        spray_cell,
                        chemical_type,
                        pass_fail,
                        mean_width,
                        mean_deflection,
                        output_image
                    ) VALUES (%s,%s,%s,%s,%s,%s,%s)
                """, (
                    now,
                    spray_cell,
                    chemical_type,
                    disp,
                    float(spray_width.mean()) if characterize else None,
                    float(deflection.mean()) if characterize else None,
                    psycopg2.Binary(image_bytes)
                ))
        
                conn.commit()
                cur.close()
                conn.close()
        
                st.success("Record saved to database!")
        
            except Exception as e:
                st.error("Database save failed.")
                st.write(str(e))
with tab2:
    
    st.header("Database History")
    
    # Connect to DB and load minimal columns first
    
    st.subheader("Filters")

    # Spray Cell Filter
    filter_cell = st.selectbox(
        "Filter by Spray Cell",
        options=["All", 1, 2, 3, 4],
        index=0
    )
    
    # Chemical Type Filter
    filter_chem = st.selectbox(
        "Filter by Chemical Type",
        options=["All", "M1b", "Glass", "MALP"],
        index=0
    )
    
    # Date Filter
    filter_date = st.date_input(
        "Filter by Date (PST)",
        value=None
    )
    
    # Convert "All" to None so it doesn't break SQL
    cell_param = None if filter_cell == "All" else filter_cell
    chem_param = None if filter_chem == "All" else filter_chem
    date_param = filter_date if filter_date else None
    
    # SQL Query
    query = """
        SELECT 
            id,
            timestamp,
            spray_cell,
            chemical_type,
            pass_fail
        FROM litmus_results
        WHERE (%s IS NULL OR spray_cell = %s)
          AND (%s IS NULL OR chemical_type = %s)
          AND (%s IS NULL OR DATE(timestamp AT TIME ZONE 'America/Los_Angeles') = %s)
        ORDER BY timestamp DESC
    """
    
    params = (
        cell_param, cell_param,
        chem_param, chem_param,
        date_param, date_param
    )
    
    # Load filtered data
    conn = get_connection()
    df = pd.read_sql(query, conn, params=params)
    conn.close()



    # Display table
    st.dataframe(df, use_container_width=True)

    st.subheader("View Record Details")

    record_id = st.number_input("Enter Record ID to view full details:", min_value=1, step=1)

    
    if st.button("Load Record"):
    
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT
                timestamp,
                spray_cell,
                chemical_type,
                pass_fail,
                mean_width,
                mean_deflection,
                output_image
            FROM litmus_results
            WHERE id = %s
        """, (record_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
    
        if row is None:
            st.error("Record not found.")
        else:
            timestamp, cell, chem, pf, mw, md, image_bytes = row
    
            st.write(f"**Timestamp:** {timestamp}")
            st.write(f"**Spray Cell:** {cell}")
            st.write(f"**Chemical Type:** {chem}")
            st.write(f"**Result:** {pf}")
            st.write(f"**Mean Width:** {mw:.3f} in")
            st.write(f"**Mean Deflection:** {md:.3f}°")
    
            if image_bytes:
                st.subheader("Stored Spray Image")
                # Convert memoryview → bytes
                img_bytes = bytes(image_bytes)
                # Now Streamlit can render it
                st.image(img_bytes, caption="Annotated Spray Pattern", width="stretch")
